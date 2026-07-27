# -*- coding: utf-8 -*-
"""Generate sample course-schedule Word docs to test each UA dictionary."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUT = Path(__file__).resolve().parent / "test_cronogramas"
OUT.mkdir(parents=True, exist_ok=True)

# Each entry: code, title, faculty, weeks of content with DOMAIN terms,
# and a final block of intentional typos (missing accents) that MUST stay red.
DOCS = [
    {
        "code": "med",
        "file": "01_cronograma_medicina.docx",
        "title": "Cronograma del curso \u2014 Medicina (prueba spellcheck med)",
        "faculty": "Facultad de Medicina / Ciencias de la Salud",
        "course": "Fisiopatolog\u00eda y Semiolog\u00eda Cl\u00ednica",
        "weeks": [
            (
                "Semana 1 \u2014 Introducci\u00f3n",
                "Farmacodin\u00e1mica, farmacocin\u00e9tica, biotransformaci\u00f3n hep\u00e1tica y aclaramiento renal. "
                "Revisi\u00f3n de homeostasia, inflamaci\u00f3n y respuesta inmunitaria.",
            ),
            (
                "Semana 2 \u2014 Cardiovascular",
                "Isquemia mioc\u00e1rdica, ateroesclerosis, miocardiopat\u00eda, arritmia y pericarditis. "
                "Ecocardiograf\u00eda, troponinas y angioplastia coronaria.",
            ),
            (
                "Semana 3 \u2014 Respiratorio",
                "Neumon\u00eda, bronquiectasia, hipoxemia, hipercapnia y ventilaci\u00f3n mec\u00e1nica. "
                "Espirometr\u00eda y gasometr\u00eda arterial.",
            ),
            (
                "Semana 4 \u2014 Laboratorio",
                "Hemograma, leucocitosis, trombocitopenia, coagulopat\u00eda y sepsis. "
                "Antibi\u00f3tico, analges\u00eda y antiinflamatorio.",
            ),
        ],
        "ok_terms": "farmacodin\u00e1mica  isquemia  miocardiopat\u00eda  hipoxemia  espirometr\u00eda  coagulopat\u00eda",
        "bad_terms": "farmacodinamica  imagenes  tecnico  inflamacion  antibiotico",
    },
    {
        "code": "odo",
        "file": "02_cronograma_odontologia.docx",
        "title": "Cronograma del curso \u2014 Odontolog\u00eda (prueba spellcheck odo)",
        "faculty": "Facultad de Odontolog\u00eda",
        "course": "Endodoncia y Periodoncia Cl\u00ednica",
        "weeks": [
            (
                "Semana 1 \u2014 Histolog\u00eda dental",
                "Amelog\u00e9nesis, dentinog\u00e9nesis, cementog\u00e9nesis, ameloblasto y odontoblasto. "
                "Desmineralizaci\u00f3n y remineralizaci\u00f3n del esmalte.",
            ),
            (
                "Semana 2 \u2014 Endodoncia",
                "Pulpectom\u00eda, pulpotom\u00eda, conductometr\u00eda, apexificaci\u00f3n y odontectom\u00eda. "
                "Radiolucidez periapical y queratoquiste odontog\u00e9nico.",
            ),
            (
                "Semana 3 \u2014 Periodoncia / implantes",
                "Periodontitis, gingivectom\u00eda, alveoloplastia, osteointegraci\u00f3n y cementoblastoma. "
                "Articulador semiajustable y oclusi\u00f3n.",
            ),
            (
                "Semana 4 \u2014 Cl\u00ednica",
                "Tartrectom\u00eda, curetaje, furcaci\u00f3n, mixoma y CBCT. "
                "Anestesia con artica\u00edna o lidoca\u00edna.",
            ),
        ],
        "ok_terms": "amelog\u00e9nesis  pulpectom\u00eda  osteointegraci\u00f3n  queratoquiste  semiajustable",
        "bad_terms": "pulpectomia  odonctectomia  extraccion  tecnico  imagenes",
    },
    {
        "code": "der",
        "file": "03_cronograma_derecho.docx",
        "title": "Cronograma del curso \u2014 Derecho (prueba spellcheck der)",
        "faculty": "Facultad de Derecho",
        "course": "Derecho Procesal y Constitucional",
        "weeks": [
            (
                "Semana 1 \u2014 Fuentes y garant\u00edas",
                "Jurisprudencia, constitucionalidad, convencionalidad y supremac\u00eda constitucional. "
                "Amparo, habeas corpus y tutela judicial efectiva.",
            ),
            (
                "Semana 2 \u2014 Proceso civil",
                "Litispendencia, litisconsorcio, casaci\u00f3n, apelaci\u00f3n y cosa juzgada. "
                "Embargo, interdicto, desahucio y ejecuci\u00f3n forzosa.",
            ),
            (
                "Semana 3 \u2014 Civil patrimonial",
                "Usufructo, fideicomiso, usucapi\u00f3n, hipoteca y anticresis. "
                "Prescripci\u00f3n extintiva y nulidad relativa.",
            ),
            (
                "Semana 4 \u2014 Penal",
                "Feminicidio, tipicidad, antijuridicidad, exequatur y ministerio p\u00fablico. "
                "Prisi\u00f3n preventiva y criterio de oportunidad.",
            ),
        ],
        "ok_terms": "amparo  casaci\u00f3n  litis  fideicomiso  usucapi\u00f3n  exequatur",
        "bad_terms": "casacion  juridico  prescripcion  apelacion  tecnico",
    },
    {
        "code": "eco",
        "file": "04_cronograma_economia.docx",
        "title": "Cronograma del curso \u2014 Ciencias Econ\u00f3micas (prueba spellcheck eco)",
        "faculty": "Facultad de Ciencias Econ\u00f3micas",
        "course": "Macroeconom\u00eda y Econometr\u00eda Aplicada",
        "weeks": [
            (
                "Semana 1 \u2014 Macro",
                "Macroeconom\u00eda, inflaci\u00f3n, deflaci\u00f3n, PIB, balanza de pagos y tipo de cambio. "
                "Pol\u00edtica monetaria y fiscalizaci\u00f3n tributaria.",
            ),
            (
                "Semana 2 \u2014 Micro / mercados",
                "Oligopolio, monopolio, elasticidad, externalidad y asimetr\u00eda de informaci\u00f3n. "
                "Arbitraje, volatilidad y diversificaci\u00f3n de portafolio.",
            ),
            (
                "Semana 3 \u2014 Finanzas",
                "Apalancamiento, amortizaci\u00f3n, depreciaci\u00f3n, liquidez y solvencia. "
                "BANGUAT, SAT, quetzal y reservas internacionales.",
            ),
            (
                "Semana 4 \u2014 Econometr\u00eda",
                "Econometr\u00eda, regresi\u00f3n, heterocedasticidad, endogeneidad y variable instrumental. "
                "Series temporales y estacionalidad.",
            ),
        ],
        "ok_terms": "oligopolio  inflaci\u00f3n  elasticidad  PIB  macroeconom\u00eda  arbitraje",
        "bad_terms": "economico  inflacion  tecnico  imagenes  credito",
    },
    {
        "code": "arq",
        "file": "05_cronograma_arquitectura.docx",
        "title": "Cronograma del curso \u2014 Arquitectura (prueba spellcheck arq)",
        "faculty": "Facultad de Arquitectura",
        "course": "Proyecto Arquitect\u00f3nico y Construcci\u00f3n",
        "weeks": [
            (
                "Semana 1 \u2014 Representaci\u00f3n",
                "Planimetr\u00eda, axonometr\u00eda, alzado, planta y corte constructivo. "
                "Croquis, volumetr\u00eda y maqueta.",
            ),
            (
                "Semana 2 \u2014 Elementos",
                "Voladizo, arquitrabe, dintel, cornisa, front\u00f3n, b\u00f3veda y c\u00fapula. "
                "Contrafuerte, pechina y z\u00f3calo.",
            ),
            (
                "Semana 3 \u2014 Estructura / normativa",
                "Cimentaci\u00f3n, encofrado, hormig\u00f3n armado y dise\u00f1o sismorresistente. "
                "COPRAQ, AGIES y urbanismo.",
            ),
            (
                "Semana 4 \u2014 Sostenibilidad",
                "Bioclim\u00e1tica, eficiencia energ\u00e9tica, patrimonio y restauraci\u00f3n. "
                "Instalaciones hidrosanitarias y el\u00e9ctricas.",
            ),
        ],
        "ok_terms": "voladizo  arquitrabe  planimetr\u00eda  bioclim\u00e1tica  sismorresistente  COPRAQ",
        "bad_terms": "construccion  arquitectonico  tecnico  imagenes  planificacion",
    },
    {
        "code": "pol",
        "file": "06_cronograma_estudios_politicos.docx",
        "title": "Cronograma del curso \u2014 Estudios Pol\u00edticos (prueba spellcheck pol)",
        "faculty": "Estudios Pol\u00edticos y Relaciones Internacionales",
        "course": "Geopol\u00edtica y Gobernanza Global",
        "weeks": [
            (
                "Semana 1 \u2014 Teor\u00eda",
                "Democracia, soberan\u00eda, gobernanza, ciudadan\u00eda y legitimidad. "
                "Geopol\u00edtica, multilateralismo y hegemon\u00eda.",
            ),
            (
                "Semana 2 \u2014 Instituciones",
                "Parlamento, plebiscito, refer\u00e9ndum, sufragio y TSE. "
                "ONU, OEA, UE y diplomacia.",
            ),
            (
                "Semana 3 \u2014 Conflictos",
                "Autoritarismo, populismo, clientelismo y polarizaci\u00f3n. "
                "Migraci\u00f3n, refugio, asilo y derechos humanos.",
            ),
            (
                "Semana 4 \u2014 Pol\u00edtica comparada",
                "Federalismo, transparencia, corrupci\u00f3n y accountability institucional. "
                "Campañas electorales y coaliciones.",
            ),
        ],
        "ok_terms": "geopol\u00edtica  soberan\u00eda  multilateralismo  plebiscito  gobernanza  ONU",
        "bad_terms": "politico  politica  eleccion  tecnico  imagenes",
    },
    {
        "code": "psi",
        "file": "07_cronograma_psicologia.docx",
        "title": "Cronograma del curso \u2014 Psicolog\u00eda (prueba spellcheck psi)",
        "faculty": "Facultad de Psicolog\u00eda",
        "course": "Psicopatolog\u00eda y Evaluaci\u00f3n Psicol\u00f3gica",
        "weeks": [
            (
                "Semana 1 \u2014 Bases",
                "Psicolog\u00eda cognitiva, neurotransmisor, psicometr\u00eda y resiliencia. "
                "Cognici\u00f3n, emocionabilidad y personalidad.",
            ),
            (
                "Semana 2 \u2014 Trastornos",
                "Psicopatolog\u00eda, depresi\u00f3n, ansiedad, esquizofrenia y trastorno bipolar. "
                "TDAH, TEA, TOC y DSM / CIE.",
            ),
            (
                "Semana 3 \u2014 Intervenci\u00f3n",
                "Psicoterapia cognitivo-conductual, psicoan\u00e1lisis y terapia sist\u00e9mica. "
                "Evaluaci\u00f3n cl\u00ednica y diagn\u00f3stico diferencial.",
            ),
            (
                "Semana 4 \u2014 M\u00e9todo",
                "Hip\u00f3tesis, muestreo, confiabilidad, validaci\u00f3n y neuropsicolog\u00eda. "
                "Variable dependiente e independiente.",
            ),
        ],
        "ok_terms": "psicopatolog\u00eda  neurotransmisor  psicometr\u00eda  resiliencia  cognitivo  DSM",
        "bad_terms": "clinico  depresion  tecnico  imagenes  psicologico",
    },
    {
        "code": "uni",
        "file": "08_cronograma_terminos_universitarios.docx",
        "title": "Cronograma del curso \u2014 T\u00e9rminos universitarios ES (prueba uni)",
        "faculty": "Universidad Francisco Marroqu\u00edn (UFM)",
        "course": "Taller de gesti\u00f3n acad\u00e9mica y syllabus",
        "weeks": [
            (
                "Semana 1 \u2014 Estructura curricular",
                "Pensum, cr\u00e9dito, semestre, prerrequisito y malla curricular. "
                "Licenciatura, maestr\u00eda y doctorado.",
            ),
            (
                "Semana 2 \u2014 Evaluaci\u00f3n",
                "R\u00fabrica, s\u00edlabus, portafolio, competencia y resultado de aprendizaje. "
                "Calificaci\u00f3n, promedio y acta.",
            ),
            (
                "Semana 3 \u2014 Administraci\u00f3n",
                "Matr\u00edcula, admisi\u00f3n PAES, acreditaci\u00f3n y convalidaci\u00f3n. "
                "Catedr\u00e1tico, decano, claustro y egresado.",
            ),
            (
                "Semana 4 \u2014 Investigaci\u00f3n",
                "Tesis, tesario, asesor\u00eda, repositorio y vinculaci\u00f3n. "
                "Cohorte, movilidad e internacionalizaci\u00f3n.",
            ),
        ],
        "ok_terms": "pensum  cr\u00e9dito  r\u00fabrica  s\u00edlabus  semestre  catedr\u00e1tico  PAES  UFM",
        "bad_terms": "credito  rubrica  tecnico  imagenes  matricula",
    },
    {
        "code": "ang",
        "file": "09_cronograma_anglicismos_campus.docx",
        "title": "Cronograma del curso \u2014 Anglicismos universitarios (prueba ang)",
        "faculty": "Campus UFM / uso acad\u00e9mico cotidiano",
        "course": "Comunicaci\u00f3n acad\u00e9mica y herramientas digitales",
        "weeks": [
            (
                "Semana 1 \u2014 Entregas",
                "Deadline del paper, extension del assignment y feedback del peer review. "
                "Checklist, handout y syllabus del workshop.",
            ),
            (
                "Semana 2 \u2014 Eventos",
                "Brainstorming, briefing, networking, coaching y mentoring. "
                "Webinar, webcast y pitch de la startup.",
            ),
            (
                "Semana 3 \u2014 Herramientas",
                "Email, online, streaming, podcast, blog y slideshow. "
                "Upload/download, meeting por Zoom/Teams y elearning en Moodle/Canvas.",
            ),
            (
                "Semana 4 \u2014 M\u00e9tricas / empleo",
                "KPI, ROI, MVP, OKR, B2B, B2C, SaaS y CRM. "
                "Internship, fellowship, freelance y soft skills.",
            ),
        ],
        "ok_terms": "feedback  deadline  paper  marketing  brainstorming  internship  B2B  SaaS  syllabus",
        "bad_terms": "tecnico  imagenes  credito  politico",
    },
]


def build_doc(spec: dict) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(spec["title"], level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(spec["faculty"])
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Curso: ").bold = True
    p.add_run(spec["course"])

    p = doc.add_paragraph()
    p.add_run("Diccionario UA a probar: ").bold = True
    p.add_run(f"dict-ua-{spec['code']}  (org.ua.dictionaries.{spec['code']}-gt)")

    p = doc.add_paragraph()
    p.add_run(
        "Instrucciones: abrir en LibreOffice Writer con idioma/locale es-GT "
        "y revisar subrayados rojos del corrector."
    ).italic = True

    doc.add_heading("1. Cronograma (t\u00e9rminos correctos de dominio)", level=1)
    for week_title, body in spec["weeks"]:
        doc.add_heading(week_title, level=2)
        doc.add_paragraph(body)

    doc.add_heading("2. Lista r\u00e1pida \u2014 deben ACEPTARSE (sin subrayado)", level=1)
    doc.add_paragraph(spec["ok_terms"])

    doc.add_heading(
        "3. Control ortogr\u00e1fico \u2014 deben RECHAZARSE (con subrayado rojo)",
        level=1,
    )
    doc.add_paragraph(
        "Estas formas van sin tilde / incorrectas a prop\u00f3sito. "
        "Si el diccionario UA las acepta, hay un problema de ortograf\u00eda prioritaria."
    )
    p = doc.add_paragraph(spec["bad_terms"])
    for run in p.runs:
        run.bold = True

    doc.add_heading("4. Nota", level=1)
    doc.add_paragraph(
        "Documento generado para pruebas del servicio spellcheck UA (es-GT). "
        "No es un syllabus oficial."
    )

    out = OUT / spec["file"]
    doc.save(out)
    return out


def main() -> int:
    paths = []
    for spec in DOCS:
        paths.append(build_doc(spec))
    readme = OUT / "README.md"
    readme.write_text(
        "# Cronogramas de prueba (spellcheck UA)\n\n"
        "Un `.docx` por diccionario. Abrir en LibreOffice con locale **es-GT**.\n\n"
        "| Archivo | Diccionario |\n"
        "|---------|-------------|\n"
        + "\n".join(
            f"| `{s['file']}` | `dict-ua-{s['code']}` |" for s in DOCS
        )
        + "\n\n"
        "**Secci\u00f3n 2:** t\u00e9rminos que deben aceptarse.\n"
        "**Secci\u00f3n 3:** faltas a prop\u00f3sito que deben seguir en rojo.\n",
        encoding="utf-8",
        newline="\n",
    )
    print(f"Escritos en: {OUT}")
    for p in paths:
        print(" ", p.name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
